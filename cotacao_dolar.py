import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from PIL import Image
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pdfkit

# Configuração do Selenium
chrome_options = Options()
arguments = ['--lang=pt-BR', '--window-size=1100,1000', '--incognito']
for argument in arguments:
    chrome_options.add_argument(argument)

chrome_options.add_experimental_option('prefs', {
    'download.prompt_for_download': False,
    'profile.default_content_setting_values.notifications': 2,
    'profile.default_content_setting_values.automatic_downloads': 1,

})

# Função para obter a cotação do dólar
def get_dollar_rate():
    url = "https://www.x-rates.com/calculator/?from=USD&to=BRL&amount=1"  # URL de exemplo
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    rate = soup.find('span', class_='ccOutputTrail').previous_sibling.text
    return rate

# Função para capturar o print da página
def capture_screenshot(url):
    driver = webdriver.Chrome(options=chrome_options)
    driver.get(url)
    screenshot = driver.get_screenshot_as_png()
    driver.quit()
    image = Image.open(BytesIO(screenshot))
    image.save('screenshot.png')

# Função para criar o relatório Word
def create_word_report(rate, date, site, screenshot_path):
    doc = Document()
    doc.add_heading('Relatório de Cotação do Dólar', 0)
    
    doc.add_paragraph(f'Cotação do Dólar: {rate}')
    doc.add_paragraph(f'Data: {date}')
    doc.add_paragraph(f'Site: {site}')
    
    doc.add_paragraph('Print do site:')
    doc.add_picture(screenshot_path, width=Inches(6))
    
    doc.save('relatorio.docx')

# Função para converter o Word em PDF
def convert_word_to_pdf(word_path, pdf_path):
    from docx2pdf import convert
    convert(word_path, pdf_path)

# Execução do script
rate = get_dollar_rate()
date = datetime.now().strftime('%Y-%m-%d')
site = "https://www.x-rates.com/calculator/?from=USD&to=BRL&amount=1"
capture_screenshot(site)
create_word_report(rate, date, site, 'screenshot.png')
convert_word_to_pdf('relatorio.docx', 'relatorio.pdf')

print(f"Cotação: {rate}, Data: {date}, Site: {site}")